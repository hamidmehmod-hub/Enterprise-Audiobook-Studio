"""
Multilingual Text & PDF-to-Speech Studio (Production Edition)
-----------------------------------------------------------------------------------
A highly modular, multithreaded application for converting massive documents into 
chapter-chunked audiobooks and video files with advanced linguistic shielding.
Includes a Sub-segmented Dual-Voice Engine, Density-Based Auto-Tagging with 
Urdu Force-Breaks, Smart Semantic Stitching, Local Database, and API Bypasses.
"""

import io
import os
import re
import gc  
import asyncio
import tempfile
import time
import zipfile
import subprocess
import hashlib
import platform
import json
import xml.sax.saxutils
from datetime import datetime
from typing import List, Tuple, Optional, Dict, Any

import numpy as np
import fitz  
import docx  
import streamlit as st
import easyocr
from PIL import Image
import google.generativeai as genai
import edge_tts
import edge_tts.communicate
from deep_translator import GoogleTranslator

# ==============================================================================
# CORE PATCH: Bypass edge-tts XML escaping to allow raw SSML injection
# ==============================================================================
if hasattr(edge_tts.communicate, 'escape'):
    edge_tts.communicate.escape = lambda x: x  


# ==============================================================================
# 1. CONFIGURATION & CONSTANTS
# ==============================================================================

LANGUAGES: Dict[str, str] = {
    "English": "en", "Urdu (اردو)": "ur", "Arabic (العربية)": "ar",
    "Persian (فارسی)": "fa", "Hindi (हिन्दी)": "hi", "Bengali (বাংলা)": "bn",
    "Spanish (Español)": "es", "French (Français)": "fr", "German (Deutsch)": "de",
    "Turkish (Türkçe)": "tr", "Russian (Русский)": "ru", "Chinese (中文)": "zh-CN",
}

DEEP_TRANS_MAP: Dict[str, str] = {
    "en": "english", "ur": "urdu", "ar": "arabic", "fa": "persian",
    "hi": "hindi", "bn": "bengali", "es": "spanish", "fr": "french",
    "de": "german", "tr": "turkish", "ru": "russian", "zh-CN": "chinese (simplified)"
}

EDGE_VOICES: Dict[str, Dict[str, Dict[str, str]]] = {
    "ur": {"Male": {"Asad Neural": "ur-PK-AsadNeural"}, "Female": {"Uzma Neural (Recommended)": "ur-PK-UzmaNeural"}},
    "en": {"Male": {"Andrew Neural": "en-US-AndrewNeural", "Brian Neural": "en-US-BrianNeural"}, "Female": {"Ava Neural": "en-US-AvaNeural", "Emma Neural": "en-US-EmmaNeural"}},
    "ar": {"Male": {"Hamed Neural": "ar-SA-HamedNeural"}, "Female": {"Zariyah Neural": "ar-SA-ZariyahNeural"}},
    "fa": {"Male": {"Farid Neural": "fa-IR-FaridNeural"}, "Female": {"Dilara Neural": "fa-IR-DilaraNeural"}},
    "hi": {"Male": {"Madhur Neural": "hi-IN-MadhurNeural"}, "Female": {"Swara Neural": "hi-IN-SwaraNeural"}},
    "bn": {"Male": {"Bashkar Neural": "bn-IN-BashkarNeural"}, "Female": {"Tanishaa Neural": "bn-IN-TanishaaNeural"}},
    "es": {"Male": {"Alvaro Neural": "es-ES-AlvaroNeural"}, "Female": {"Elvira Neural": "es-ES-ElviraNeural"}},
    "fr": {"Male": {"Henri Neural": "fr-FR-HenriNeural"}, "Female": {"Denise Neural": "fr-FR-DeniseNeural"}},
    "de": {"Male": {"Conrad Neural": "de-DE-ConradNeural"}, "Female": {"Katja Neural": "de-DE-KatjaNeural"}},
    "tr": {"Male": {"Ahmet Neural": "tr-TR-AhmetNeural"}, "Female": {"Emel Neural": "tr-TR-EmelNeural"}},
    "ru": {"Male": {"Dmitry Neural": "ru-RU-DmitryNeural"}, "Female": {"Svetlana Neural": "ru-RU-SvetlanaNeural"}},
    "zh-CN": {"Male": {"Yunxi Neural": "zh-CN-YunxiNeural"}, "Female": {"Xiaoxiao Neural": "zh-CN-XiaoxiaoNeural"}}
}


# ==============================================================================
# 2. SSML DIRECTOR'S CUT & DUAL-VOICE ENGINE
# ==============================================================================

def apply_ssml_director_cuts(text: str) -> str:
    """Parses user-friendly director tags into raw SSML XML without touching [sec] tags."""
    text = xml.sax.saxutils.escape(text)
    
    text = re.sub(r'\[pause\s+([a-zA-Z0-9.]+)\]', r'<break time="\1"/>', text, flags=re.IGNORECASE)
    text = re.sub(r'\[emphasis\](.*?)\[/emphasis\]', r'<emphasis level="strong">\1</emphasis>', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'\[whisper\](.*?)\[/whisper\]', r'<prosody volume="x-soft" rate="-15%" pitch="+10%">\1</prosody>', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'\[shout\](.*?)\[/shout\]', r'<prosody volume="x-loud" rate="+10%" pitch="+10%">\1</prosody>', text, flags=re.IGNORECASE | re.DOTALL)
        
    return text

def strip_director_tags(text: str) -> str:
    """Removes director tags for clean SRT subtitle generation."""
    text = re.sub(r'\[pause\s+[a-zA-Z0-9.]+\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[/?emphasis\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[/?whisper\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[/?shout\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\[/?sec\]', '', text, flags=re.IGNORECASE)
    return re.sub(r'\s+', ' ', text).strip()

def auto_tag_arabic_script(text: str) -> str:
    """Intelligently wraps Arabic verses in [sec] tags using Density Engine & Waqf Bridge."""
    text = re.sub(r'\[/?sec\]', '', text, flags=re.IGNORECASE) 
    
    tokens = re.split(r'(\s+)', text)
    tagged_tokens = []
    in_sec = False
    
    arabic_connectors = {
        'و', 'ف', 'ب', 'ك', 'ل', 'من', 'في', 'عن', 'على', 'الى', 'إلى', 
        'هو', 'هي', 'قد', 'لم', 'لن', 'ثم', 'الا', 'إلا'
    }
    
    urdu_force_breaks = {
        'اس', 'ان', 'کہ', 'یہ', 'وہ', 'اور', 'اُس', 'اِس', 'اِن', 'اُن', 
        'کی', 'کے', 'کو', 'سے', 'میں', 'نے', 'ہیں', 'ہے', 'تھا', 'تھی', 'تھے', 'کا', 'جو', 'پر'
    }
    
    for token in tokens:
        if not token.strip(): 
            tagged_tokens.append(token)
            continue
            
        clean_word = re.sub(r'[\u064B-\u065F\u06D6-\u06DC\u0670]', '', token).strip()
        diacritic_count = len(re.findall(r'[\u064B-\u065F\u06D6-\u06DC\u0670]', token))
        has_brackets = bool(re.search(r'[﴿﴾]', token))
        
        is_punct = bool(re.match(r'^[\W_]+$', clean_word)) if clean_word else False
        is_only_diacritics = len(clean_word) == 0 and len(token.strip()) > 0
        
        if not in_sec:
            if has_brackets or diacritic_count >= 2:
                in_sec = True
                tagged_tokens.append("[sec]")
            tagged_tokens.append(token)
        else:
            keep_on = has_brackets or is_punct or is_only_diacritics or diacritic_count >= 2 or clean_word in arabic_connectors
            
            if clean_word in urdu_force_breaks:
                keep_on = False
                
            if not keep_on:
                in_sec = False
                tagged_tokens.append("[/sec]")
            tagged_tokens.append(token)
            
    if in_sec:
        tagged_tokens.append("[/sec]")
        
    result = "".join(tagged_tokens)
    result = re.sub(r'\[sec\](\s*)\[/sec\]', r'\1', result)
    return result


# ==============================================================================
# 3. TEXT PROCESSING & CHUNKING ENGINE
# ==============================================================================

def chunk_by_words(text: str, max_words: int = 3000, title_prefix: str = "Part") -> List[Tuple[str, str]]:
    paragraphs = text.split('\n')
    chunks, current_chunk = [], []
    current_words, part_num = 0, 1
    
    for p in paragraphs:
        p_words = len(p.split())
        if current_words + p_words > max_words and current_words > 0:
            chunks.append((f"{title_prefix} {part_num}", "\n".join(current_chunk)))
            part_num += 1
            current_chunk = [p]
            current_words = p_words
        else:
            current_chunk.append(p)
            current_words += p_words
            
    if current_chunk:
        chunks.append((f"{title_prefix} {part_num}", "\n".join(current_chunk)))
    return chunks

def hybrid_chapter_splitter(text: str, max_words: int = 3000) -> List[Tuple[str, str]]:
    if not text.strip(): return []
        
    marker_pattern = re.compile(r'^\s*(chapter|باب|جلد|الفصل|অধ্যায়|भाग|глава|bölüm|capítulo|chapitre|kapitel|第).*', re.IGNORECASE | re.UNICODE)
    lines = text.split('\n')
    chapters, current_lines = [], []
    current_title = "Prologue"
    marker_found = False
    
    for line in lines:
        clean_line = line.strip()
        if marker_pattern.match(clean_line) and len(clean_line.split()) <= 10:
            marker_found = True
            if "".join(current_lines).strip():
                chapters.append((current_title, "\n".join(current_lines)))
            current_title = re.sub(r'[\\/*?:"<>|]', "", clean_line)[:60]
            current_lines = []
        else:
            current_lines.append(line)
            
    if "".join(current_lines).strip():
        chapters.append((current_title, "\n".join(current_lines)))
        
    if not marker_found:
        return chunk_by_words(text, max_words)
        
    refined_chapters = []
    for title, content in chapters:
        if len(content.split()) > max_words * 1.5:
            sub_chunks = chunk_by_words(content, max_words, title_prefix=f"{title} - Segment")
            refined_chapters.extend(sub_chunks)
        else:
            refined_chapters.append((title, content))
            
    return refined_chapters

def restore_quranic_verses(text: str) -> str:
    if not text: return ""
    verse_repairs = {
        r'لیلے\s+لس\s+خلق\s+سب\s+سموت.*?(?=بخاری|\n|$)': 
            'اللَّهُ الَّذِي خَلَقَ سَبْعَ سَمَاوَاتٍ وَمِنَ الْأَرْضِ مِثْلَهُنَّ يَتَنَزَّلُ الْأَمْرُ بَيْنَهُنَّ لِتَعْلَمُوا أَنَّ اللَّهَ عَلَىٰ كُلِّ شَيْءٍ قَدِيرٌ وَأَنَّ اللَّهَ قَدْ أَحَاطَ بِكُلِّ شَيْءٍ عِلْمًا۔ ',
    }
    for pattern, replacement in verse_repairs.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE | re.DOTALL)
    return text

def advanced_urdu_normalizer(text: str) -> str:
    if not text: return ""
    
    text = re.sub(r'(م[\u064B-\u065F]*ل[\u064B-\u065F]*[یي]?|صلیٰ|مئی ایم|صلی السلام|صلی علیم|مئی)\s+(صلی اللہ علیہ وآلہ وسلم)', r'\2', text)
    text = re.sub(r'(آنحضرت|رسول\s*اللہ|نبی\s*کریم)\s+(صلیٰ?\s*ا\s*ہم|صلی\s*الہ\s*ہم|صلی\s*السلام|مئی\s*ایم|صلی\s*ایم|مَلَّی\s*اَللہ|مَلَّی\s*الم|صلی\s*علیم|علی\s*الم)', r'\1 صلی اللہ علیہ وآلہ وسلم', text)
    text = re.sub(r'(آنحضرت|رسول\s*اللہ|نبی\s*کریم)\s+(صلیٰ|مَلِ|مَلَّی|ملی|مئی)(?!\w)', r'\1 صلی اللہ علیہ وآلہ وسلم', text)
    
    pbuh_super_regex = (
        r'[صم][\u064B-\u065F]*[لئ][\u064B-\u065F]*[یيےى]?\s+[\u064B-\u065F]*'
        r'(?:اللہ|اﷲ|الہ)'
        r'(?:\s+[\u064B-\u065F]*(?:علیہ|علیم|علہ|علی))?'
        r'(?:\s+[\u064B-\u065F]*(?:وآلہ|وآلہ))?'
        r'(?:\s+[\u064B-\u065F]*(?:وسلم|الم|ہم|وسلمہ))?'
    )
    text = re.sub(pbuh_super_regex, ' صلی اللہ علیہ وآلہ وسلم ', text)
    
    text = re.sub(r'(صلی اللہ علیہ وآلہ وسلم)\s*[\u064B-\u065F]*\s*(?:الْوَیَّمْ|الویّم|الْوَیَّم)[\u064B-\u065F]*', r'\1', text)
    
    text = re.sub(r'ع[\u064B-\u065F]*ل[\u064B-\u065F]*[یيے][\u064B-\u065F]*[ہه][\u064B-\u065F]*م[\u064B-\u065F]*\s*ا?ل[\u064B-\u065F]*س[\u064B-\u065F]*ل[\u064B-\u065F]*ا[\u064B-\u065F]*م[\u064B-\u065F]*', 'علیہم السلام', text)
    text = re.sub(r'ع[\u064B-\u065F]*ل[\u064B-\u065F]*[یيے][\u064B-\u065F]*[ہه][\u064B-\u065F]*\s*ا?ل[\u064B-\u065F]*س[\u064B-\u065F]*ل[\u064B-\u065F]*ا[\u064B-\u065F]*م[\u064B-\u065F]*', 'علیہ السلام', text)
    text = re.sub(r'ر[\u064B-\u065F]*ح[\u064B-\u065F]*م[\u064B-\u065F]*[ہهةت][\u064B-\u065F]*', 'رحمہ', text)
    
    honorifics_map = {
        r'ﷺ': ' صلی اللہ علیہ وآلہ وسلم ', 
        r'ؓ': ' رضی اللہ تعالیٰ عنہ ',
        r'ؒ': ' رحمۃ اللہ علیہ ', 
        r'ؑ': ' علیہ السلام ', 
        r'ﷻ': ' جل جلالہ ',
        r'آںحضرت': 'آنحضرت', 
        r'اﷲ': 'اللہ',
        r'رضی\s*اللہ\s*(?:عنہا|عنھا)': ' رضی اللہ عنہا ', 
        r'رضی\s*اللہ\s*عنہ(?!\w)': ' رضی اللہ عنہ ', 
        r'رضی\s*اللہ\s*نہم': ' رضی اللہ عنہم ',
    }
    for pattern, replacement in honorifics_map.items():
        text = re.sub(pattern, replacement, text)
        
    text = re.sub(r'(صلی اللہ علیہ وآلہ وسلم\s*)+', 'صلی اللہ علیہ وآلہ وسلم ', text)
    return text.replace('\u200c', ' ').replace('\u200d', '')

def clean_extracted_text(text: str) -> str:
    if not text: return ""
    text = restore_quranic_verses(text)
    
    text = re.sub(r'^\s*[\*\•\>]\s*', '', text, flags=re.MULTILINE)
    
    filtered_lines = []
    meta_keywords = ['check arabic', 'header:', 'image has:', 'left)', 'arabic quote:', 'verbatim', 'end_of_page']
    for line in text.split('\n'):
        clean_line = line.strip()
        if not clean_line or any(kw in clean_line.lower() for kw in meta_keywords): 
            continue
        filtered_lines.append(clean_line)

    stitched_lines = []
    current_paragraph = ""
    
    for line in filtered_lines:
        if not current_paragraph:
            current_paragraph = line
        else:
            last_char = current_paragraph[-1]
            is_end_of_sentence = last_char in ['۔', '؟', '!', '.', '”', '"', '﴾']
            is_short_line = len(current_paragraph.split()) <= 8
            
            if is_end_of_sentence or is_short_line:
                stitched_lines.append(current_paragraph)
                current_paragraph = line
            else:
                current_paragraph += " " + line
                
    if current_paragraph:
        stitched_lines.append(current_paragraph)

    unique_blocks, seen = [], set()
    for block in stitched_lines:
        fingerprint = re.sub(r'[\s\u064B-\u0652\u0670]', '', block)[:40]
        if len(fingerprint) < 20 or fingerprint not in seen:
            seen.add(fingerprint)
            unique_blocks.append(block)
            
    text = "\n\n".join(unique_blocks)
    text = re.sub(r'\b(\w+)(?:\s+\1)+\b', r'\1', text, flags=re.UNICODE) 
    
    return advanced_urdu_normalizer(re.sub(r'[ \t]+', ' ', text))

def reverse_line_words(text: str) -> str:
    return "\n".join([" ".join(reversed(line.split())) for line in text.split('\n')])


# ==============================================================================
# 4. TRANSLATION ENGINE
# ==============================================================================

def translate_text_safely(text: str, target_lang_string: str) -> str:
    if not text.strip(): return ""
    
    shielded_verse = 'اللَّهُ الَّذِي خَلَقَ سَبْعَ سَمَاوَاتٍ وَمِنَ الْأَرْضِ مِثْلَهُنَّ يَتَنَزَّلُ الْأَمْرُ بَيْنَهُنَّ لِتَعْلَمُوا أَنَّ اللَّهَ عَلَىٰ كُلِّ شَيْءٍ قَدِيرٌ وَأَنَّ اللَّهَ قَدْ أَحَاطَ بِكُلِّ شَيْءٍ عِلْمًا'
    token = " XXQURAN999XX "
    text = text.replace(shielded_verse, token)

    translator = GoogleTranslator(source='auto', target=target_lang_string)
    translated_paragraphs, current_chunk = [], ""
    
    def fetch_with_retry(chunk: str, retries: int = 3) -> str:
        for attempt in range(retries):
            try:
                time.sleep(3) 
                return translator.translate(chunk)
            except Exception as e:
                if attempt == retries - 1: raise Exception(f"API Blocked: {str(e)}")
                time.sleep(5) 

    for p in text.split('\n'):
        if len(current_chunk) + len(p) < 3500:
            current_chunk += p + "\n"
        else:
            if current_chunk.strip(): translated_paragraphs.append(fetch_with_retry(current_chunk))
            current_chunk = p + "\n"
    if current_chunk.strip():
        translated_paragraphs.append(fetch_with_retry(current_chunk))
        
    final_translation = "\n".join(translated_paragraphs)
    
    target_lower = target_lang_string.lower()
    if target_lower == "turkish": final_translation = re.sub(r'(?i)Gaspçı', 'İşgalci', final_translation)
    elif target_lower == "persian": final_translation = re.sub(r'\bاور\b', 'و', final_translation)
    elif target_lower == "spanish": final_translation = re.sub(r'(?i)\bHaya\b', 'Alá', final_translation)
    
    return re.sub(r'(?i)XXQURAN999XX', shielded_verse, final_translation)


# ==============================================================================
# 5. EXTRACTION ENGINE
# ==============================================================================

def extract_text_from_docx(file_bytes: bytes) -> str:
    return "\n".join([para.text for para in docx.Document(io.BytesIO(file_bytes)).paragraphs if para.text.strip()])

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode('utf-8', errors='ignore')

@st.cache_resource
def load_ocr_reader(lang_code: str):
    target_langs = [lang_code, "en"]
    if lang_code in ["ur", "fa"] and "ar" not in target_langs: target_langs.append("ar")
    if lang_code in ["hi", "bn", "es", "fr", "de", "tr", "ru", "zh-CN"] and lang_code not in target_langs:
        target_langs.append(lang_code)
    if "zh-CN" in target_langs:
        target_langs.remove("zh-CN")
        target_langs.append("ch_sim")
    return easyocr.Reader(target_langs, gpu=True)

def extract_text_ocr_local(pdf_bytes: bytes, lang_code: str, start_page: int, end_page: int, progress_bar: Any) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    reader = load_ocr_reader(lang_code)
    full_text, actual_end = [], min(end_page, len(doc))
    
    for idx, page_num in enumerate(range(start_page - 1, actual_end)):
        pix = doc[page_num].get_pixmap(dpi=150)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
        if pix.n == 4: img = img[:, :, :3]
        
        full_text.append(" ".join(reader.readtext(img, detail=0)))
        if progress_bar: progress_bar.progress((idx + 1) / (actual_end - (start_page - 1)))
        
        del pix
        del img
        gc.collect()

    doc.close() 
    return clean_extracted_text("\n\n".join(full_text))

def extract_text_with_vision_ai(pdf_bytes: bytes, api_key: str, start_page: int, end_page: int, progress_bar: Any, file_index: int=0, total_files: int=1) -> str:
    genai.configure(api_key=api_key)
    try:
        models = [m.name.replace("models/", "") for m in genai.list_models() if "generateContent" in m.supported_generation_methods]
        flash_models = sorted([m for m in models if "flash" in m and "preview" not in m], reverse=True)
        FAST_MODEL = flash_models[0] if flash_models else "gemini-1.5-flash-latest"
    except Exception:
        FAST_MODEL = "gemini-1.5-flash-latest"

    safety_settings = {
        'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE',
        'HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE',
        'HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE',
        'HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'
    }

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    actual_end = min(end_page, len(doc))
    total_pages = actual_end - (start_page - 1)
    
    prompt = (
        "You are an expert OCR transcriber.\n"
        "Please transcribe the ENTIRE text visible in this image from top to bottom exactly as written.\n"
        "CRITICAL RULE: If you see Arabic verses or quotes embedded in Urdu, transcribe them using EXACT Arabic spelling with ALL original diacritics and harakat intact. If the text is English, transcribe it purely in English.\n"
        "Output only the raw body text. Do not include any comments or introductions."
    )
    model = genai.GenerativeModel(FAST_MODEL, generation_config=genai.types.GenerationConfig(temperature=0.3))
    transcribed_pages = []
    
    for idx, page_num in enumerate(range(start_page - 1, actual_end)):
        pix = doc[page_num].get_pixmap(dpi=150)
        img_bytes = pix.tobytes("jpeg", 85)
        if idx > 0: time.sleep(4) 
        
        try:
            res = model.generate_content([prompt, Image.open(io.BytesIO(img_bytes))], safety_settings=safety_settings)
            try:
                transcribed_pages.append(res.text.strip() if res.text else f"\n[Page {page_num + 1} API Error: Empty]\n")
            except ValueError:
                transcribed_pages.append(f"\n[Page {page_num + 1} Error: Gemini blocked the content due to filters.]\n")
        except Exception as e:
            transcribed_pages.append(f"\n[Page {page_num + 1} API Error: {str(e)}]\n")
            
        if progress_bar: progress_bar.progress((file_index / total_files) + (((idx + 1) / total_pages) * (1 / total_files)))
        
        del pix
        gc.collect()

    doc.close() 
    return clean_extracted_text("\n\n".join(transcribed_pages))


# ==============================================================================
# 6. MEDIA GENERATION ENGINE
# ==============================================================================

def generate_docx(text: str) -> bytes:
    doc = docx.Document()
    doc.add_heading('Generated Script', 0)
    for p in [p.strip() for p in text.split('\n') if p.strip()]: doc.add_paragraph(p)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

def create_zip_archive(chapter_data_list: List[Dict[str, Any]], full_docx_bytes: Optional[bytes] = None) -> bytes:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        if full_docx_bytes: zip_file.writestr("Full_Script_Document.docx", full_docx_bytes)
            
        for i, data in enumerate(chapter_data_list):
            title = f"{i+1:03d}_{data['title'].replace(' ', '_')}"
            if data.get('audio_bytes'): zip_file.writestr(f"{title}/{title}.mp3", data['audio_bytes'])
            if data.get('srt_text'): zip_file.writestr(f"{title}/{title}.srt", data['srt_text'].encode('utf-8'))
            if data.get('script_text'): zip_file.writestr(f"{title}/{title}.txt", data['script_text'].encode('utf-8'))
            if data.get('video_bytes'): zip_file.writestr(f"{title}/{title}.mp4", data['video_bytes'])
                
    return zip_buffer.getvalue()

def create_master_zip_archive(history: List[Dict[str, Any]]) -> bytes:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for item in history:
            if item.get('zip_bytes'):
                zip_file.writestr(item['filename'], item['zip_bytes'])
    return zip_buffer.getvalue()

def generate_mp4_video(audio_bytes: bytes, srt_text: str, total_duration: float, progress_bar: Any) -> bytes:
    with tempfile.TemporaryDirectory() as temp_dir:
        audio_path, srt_path = os.path.join(temp_dir, "a.mp3"), os.path.join(temp_dir, "s.srt")
        img_path, out_path = os.path.join(temp_dir, "b.jpg"), os.path.join(temp_dir, "o.mp4")

        with open(audio_path, "wb") as f: f.write(audio_bytes)
        with open(srt_path, "w", encoding="utf-8") as f: f.write(srt_text)
        Image.new('RGB', (1280, 720), color=(30, 30, 30)).save(img_path)

        cmd = [
            "ffmpeg", "-y", "-loop", "1", "-framerate", "15", "-i", "b.jpg", "-i", "a.mp3",
            "-vf", "subtitles=s.srt:force_style='FontSize=24,PrimaryColour=&H00FFFFFF,Outline=1,Shadow=1,MarginV=50'",
            "-c:v", "libx264", "-preset", "ultrafast", "-tune", "stillimage", "-threads", "0",
            "-c:a", "copy", "-shortest", "-pix_fmt", "yuv420p", "o.mp4"
        ]

        try:
            proc = subprocess.Popen(cmd, cwd=temp_dir, stdout=subprocess.PIPE, stderr=subprocess.PIPE, universal_newlines=True)
            for line in proc.stderr:
                match = re.search(r"time=(\d+):(\d+):(\d+\.\d+)", line)
                if match and total_duration > 0 and progress_bar:
                    sec = int(match.group(1))*3600 + int(match.group(2))*60 + float(match.group(3))
                    progress_bar.progress(min(sec / total_duration, 1.0), text=f"Rendering video... {int((sec/total_duration)*100)}%")
            proc.wait()
            if proc.returncode != 0: raise Exception("FFmpeg failed.")
            with open(out_path, "rb") as f: return f.read()
        except FileNotFoundError: raise Exception("FFmpeg is missing from system PATH.")

async def _async_generate_parallel_audio(text: str, voice_code: str, sec_voice_code: str, rate: str, pitch: str, progress_bar: Any = None) -> Tuple[bytes, str, float]:
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 2]
    if not paragraphs: return b"", "", 0.0
        
    results = [(b"", 0.0)] * len(paragraphs)
    sem = asyncio.Semaphore(3) 
    
    async def fetch_chunk(idx: int, p: str, max_retries: int = 4):
        parts = re.split(r'(\[sec\].*?\[/sec\])', p, flags=re.IGNORECASE | re.DOTALL)
        chunk_audio = b""
        chunk_duration = 0.0
        
        async with sem:
            for part in parts:
                if not part.strip(): continue
                
                is_sec = part.lower().startswith('[sec]') and part.lower().endswith('[/sec]')
                if is_sec:
                    target_text = part[5:-6]
                    target_voice = sec_voice_code if sec_voice_code else voice_code
                else:
                    target_text = part
                    target_voice = voice_code
                    
                if not target_text.strip(): continue
                
                p_ssml = apply_ssml_director_cuts(target_text)
                
                part_audio = b""
                part_dur = 0.0
                success = False
                
                for attempt in range(max_retries):
                    try:
                        await asyncio.sleep(0.5) 
                        async for api_chunk in edge_tts.Communicate(p_ssml, target_voice, rate=rate, pitch=pitch).stream():
                            if api_chunk["type"] == "audio": 
                                part_audio += api_chunk["data"]
                            elif api_chunk["type"] == "WordBoundary":
                                part_dur = max(part_dur, (api_chunk["offset"] + api_chunk["duration"]) / 10000000.0)
                        
                        if len(part_audio) > 0:
                            if part_dur == 0.0: part_dur = len(part_audio) / 6000.0 
                            chunk_audio += part_audio
                            chunk_duration += part_dur + 0.1
                            success = True
                            break
                        else:
                            raise Exception("Empty audio received from server.")
                            
                    except Exception as e: 
                        print(f"TTS Error in Chunk {idx} (Attempt {attempt+1}): {e}")
                        if attempt < max_retries - 1:
                            await asyncio.sleep(2 ** attempt) 
                            
                if not success:
                    print(f"\n=======================================================")
                    print(f"CRITICAL FAILURE: Could not render part of chunk {idx}")
                    print(f"--- TEXT CONTENT OF FAILING CHUNK ---")
                    print(p_ssml)
                    print(f"=======================================================\n")
                    return idx, b"", 0.0
            
            return idx, chunk_audio, chunk_duration
            
    for coro in asyncio.as_completed([fetch_chunk(i, p) for i, p in enumerate(paragraphs)]):
        idx, a_data, dur = await coro
        results[idx] = (a_data, dur)
            
    final_audio, srt_content, curr_time = b"", "", 0.0
    def f_time(s: float) -> str: return f"{int(s//3600):02d}:{int((s%3600)//60):02d}:{int(s%60):02d},{int((s%1)*1000):03d}"

    for i, (a_data, dur) in enumerate(results):
        if not a_data: continue
        final_audio += a_data
        
        clean_srt_text = strip_director_tags(paragraphs[i])
        
        srt_content += f"{i + 1}\n{f_time(curr_time)} --> {f_time(curr_time + dur)}\n{clean_srt_text}\n\n"
        curr_time += dur
        
    return final_audio, srt_content, curr_time


# ==============================================================================
# 7. DATABASE & METADATA MANAGEMENT
# ==============================================================================

def update_project_metadata(workspace_dir: str, doc_name: str, lang: str, voice: str, total_chap: int, comp_chap: int):
    meta_path = os.path.join(workspace_dir, "metadata.json")
    metadata = {
        "doc_name": doc_name,
        "language": lang,
        "voice": voice,
        "total_chapters": total_chap,
        "completed_chapters": comp_chap,
        "last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=4)

def package_folder_to_zip(folder_path: str) -> bytes:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if file != "metadata.json" and file != "master_script.txt" and file != "chunk_hash.txt":
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, folder_path)
                    zip_file.write(file_path, arcname)
    return zip_buffer.getvalue()


# ==============================================================================
# 8. STREAMLIT UI COMPONENTS
# ==============================================================================

def initialize_session_state():
    if "audio_history" not in st.session_state: st.session_state.audio_history = []
    if "trans_error" not in st.session_state: st.session_state.trans_error = None
    if "pasted_text" not in st.session_state: st.session_state.pasted_text = ""
    if "pdf_extracted_text" not in st.session_state: st.session_state.pdf_extracted_text = None
    if "source_doc_name" not in st.session_state: st.session_state.source_doc_name = "Pasted_Text"

def render_sidebar() -> Tuple[str, str, str, str, str, bool, str, str, Optional[str]]:
    with st.sidebar:
        st.header("⚙️ Primary Voice (Main Text)")
        lang_lbl = st.selectbox("Target Language", options=list(LANGUAGES.keys()), index=0)
        lang_cd = LANGUAGES[lang_lbl]

        genders = list(EDGE_VOICES.get(lang_cd, {}).keys())
        default_gender_idx = genders.index("Female") if "Female" in genders else 0
        gender = st.radio("Voice Gender", options=genders, index=default_gender_idx, horizontal=True)

        personas = EDGE_VOICES.get(lang_cd, {}).get(gender, {})
        if personas:
            persona_lbl = st.selectbox("Voice Persona", options=list(personas.keys()), index=0)
            voice_cd = personas[persona_lbl]
        else:
            st.warning("No voices available.")
            voice_cd, persona_lbl = None, None

        st.markdown("---")
        
        st.header("🌐 Secondary Voice (Quotes/Arabic)")
        enable_sec_voice = st.checkbox("Enable Dual-Voice Engine", value=False)
        sec_voice_cd = None
        if enable_sec_voice:
            sec_lang_lbl = st.selectbox("Secondary Language", options=list(LANGUAGES.keys()), index=2) 
            sec_lang_cd = LANGUAGES[sec_lang_lbl]
            
            sec_genders = list(EDGE_VOICES.get(sec_lang_cd, {}).keys())
            sec_gender = st.radio("Sec. Voice Gender", options=sec_genders, horizontal=True, key="sec_gender")
            
            sec_personas = EDGE_VOICES.get(sec_lang_cd, {}).get(sec_gender, {})
            if sec_personas:
                sec_persona_lbl = st.selectbox("Sec. Voice Persona", options=list(sec_personas.keys()), index=0, key="sec_persona")
                sec_voice_cd = sec_personas[sec_persona_lbl]
            else:
                st.warning("No secondary voices available.")
        
        st.markdown("---")
        st.header("🎛️ Global Audio Controls")
        rate = st.slider("Speech Speed (%)", min_value=-50, max_value=50, value=0, step=5)
        pitch = st.slider("Voice Pitch (Hz)", min_value=-50, max_value=50, value=0, step=5)
        
        st.markdown("---")
        st.header("🎥 Video Generation")
        exp_mp4 = st.checkbox("🎬 Export MP4 (Requires FFmpeg)", value=False)
        
    return lang_cd, lang_lbl, voice_cd, f"{rate:+d}%", f"{pitch:+d}Hz", exp_mp4, gender, persona_lbl, sec_voice_cd

def render_error_dashboard():
    if st.session_state.trans_error:
        st.error(st.session_state.trans_error)
        if st.button("Dismiss Error & Reset"):
            st.session_state.trans_error = None
            st.rerun()
        st.divider()

def render_script_tools(state_key: str, lang_cd: str, lang_lbl: str):
    col1, col2, col3, col4 = st.columns(4)
    text_val = st.session_state[state_key]
    
    with col1:
        if st.button("🔄 Reverse Words", width="stretch"):
            st.session_state[state_key] = reverse_line_words(text_val)
            st.rerun()
    with col2:
        if st.button("🧹 Auto-Clean", width="stretch"):
            st.session_state[state_key] = clean_extracted_text(text_val)
            st.rerun()
    with col3:
        if st.button(f"🌐 Translate to {lang_lbl}", width="stretch"):
            with st.spinner("Translating safely..."):
                try:
                    dt_str = DEEP_TRANS_MAP.get(lang_cd, "english")
                    st.session_state[state_key] = translate_text_safely(text_val, dt_str)
                    st.rerun()
                except Exception as e:
                    st.session_state.trans_error = str(e)
                    st.rerun()
    with col4:
        if st.button("🪄 Auto-Tag Arabic", width="stretch"):
            st.session_state[state_key] = auto_tag_arabic_script(text_val)
            st.rerun()

    with st.expander("🎬 Director's Cut: SSML Audio Effects (Cheat Sheet)", expanded=False):
        st.markdown("""
        You can manually inject audio effects directly into your script! Type these tags anywhere in the text below:
        *   **Secondary Voice:** `[sec] text [/sec]` *(Switches to your selected secondary voice for Arabic/Quotes)*
        *   **Pause/Breathe:** `[pause 2s]` or `[pause 500ms]` *(Adds silence)*
        *   **Whisper:** `[whisper] text [/whisper]` *(Softens volume, lowers rate)*
        *   **Shout:** `[shout] text [/shout]` *(Increases volume and pitch)*
        *   **Emphasis:** `[emphasis] text [/emphasis]` *(Puts hard emphasis on the words)*
        
        *Note: Do not span tags across multiple paragraphs. Close the tag before hitting Enter.*
        """)

    st.session_state[state_key] = st.text_area("Master Script Editor:", value=st.session_state[state_key], height=400)
    return st.session_state[state_key]

def handle_generation(active_text: str, lang_cd: str, lang_lbl: str, voice_cd: str, rate: str, pitch: str, exp_mp4: bool, gender: str, persona: str, sec_voice_cd: str):
    if not active_text.strip(): return
    st.divider()
    
    if not voice_cd:
        st.warning("Please select a Primary Voice Persona from the sidebar.")
        return

    final_script = clean_extracted_text(active_text)
    
    doc_name_clean = re.sub(r'[\\/*?:"<>|]', '', st.session_state.source_doc_name)[:50]
    folder_name = f"{doc_name_clean}_{voice_cd}_{rate.replace('%','')}_{pitch.replace('Hz','')}"
    workspace_dir = os.path.join(os.getcwd(), "audiobook_workspace", folder_name)
    os.makedirs(workspace_dir, exist_ok=True)
    
    dual_status = "🟢 Dual-Voice Active" if sec_voice_cd else "⚪ Single Voice"
    st.write(f"**Total Words:** {len(active_text.split()):,} | **Output:** {lang_lbl} ({gender}) | {dual_status}")
    
    col_info, col_btn = st.columns([4, 1])
    with col_info:
        st.info(f"📂 **Active Workspace (Auto-Save):**\n`{workspace_dir}`")
    with col_btn:
        if st.button("📁 Open Folder", width="stretch"):
            try:
                if platform.system() == "Windows": os.startfile(workspace_dir)
                elif platform.system() == "Darwin": subprocess.Popen(["open", workspace_dir])
                else: subprocess.Popen(["xdg-open", workspace_dir])
            except Exception as e: st.error(f"Error opening folder: {e}")

    if st.button("🎙️ Generate Audiobook (Auto-Chaptering)", type="primary", width="stretch"):
        master_script_path = os.path.join(workspace_dir, "master_script.txt")
        with open(master_script_path, "w", encoding="utf-8") as f:
            f.write(final_script)
            
        full_docx = generate_docx(strip_director_tags(final_script))
        chapters = hybrid_chapter_splitter(final_script, max_words=3000)
        
        update_project_metadata(workspace_dir, doc_name_clean, lang_lbl, f"{gender} - {persona}", len(chapters), 0)
        
        progress_bar = st.progress(0.0, text="Initializing Smart Resume Engine...")
        chapter_outputs = []
        completed_count = 0
        
        try:
            for idx, (title, content) in enumerate(chapters):
                safe_title = f"{idx+1:03d}_{re.sub(r'[\\\\/*?:\"<>|]', '', title.replace(' ', '_'))}"
                chapter_dir = os.path.join(workspace_dir, safe_title)
                os.makedirs(chapter_dir, exist_ok=True)
                
                mp3_path = os.path.join(chapter_dir, f"{safe_title}.mp3")
                srt_path = os.path.join(chapter_dir, f"{safe_title}.srt")
                txt_path = os.path.join(chapter_dir, f"{safe_title}.txt")
                mp4_path = os.path.join(chapter_dir, f"{safe_title}.mp4")
                
                cache_str = f"{content}_{sec_voice_cd}"
                chunk_hash = hashlib.md5(cache_str.encode('utf-8')).hexdigest()[:8]
                hash_path = os.path.join(chapter_dir, "chunk_hash.txt")
                
                needs_processing = True
                if os.path.exists(mp3_path) and os.path.exists(srt_path):
                    if os.path.exists(hash_path):
                        with open(hash_path, "r") as f:
                            saved_hash = f.read().strip()
                        if saved_hash == chunk_hash:
                            if exp_mp4:
                                if os.path.exists(mp4_path): needs_processing = False
                            else:
                                needs_processing = False
                
                if not needs_processing:
                    progress_bar.progress(idx / len(chapters), text=f"Restoring from disk: {title} ({idx+1}/{len(chapters)})")
                    with open(mp3_path, "rb") as f: audio_b = f.read()
                    with open(srt_path, "r", encoding="utf-8") as f: srt_b = f.read()
                    vid_b = None
                    if exp_mp4:
                        with open(mp4_path, "rb") as f: vid_b = f.read()
                        
                    chapter_outputs.append({"title": title, "audio_bytes": audio_b, "srt_text": srt_b, "script_text": content, "video_bytes": vid_b})
                    
                else:
                    progress_bar.progress(idx / len(chapters), text=f"Audio: {title} ({idx+1}/{len(chapters)})")
                    audio_b, srt_b, dur = asyncio.run(_async_generate_parallel_audio(content, voice_cd, sec_voice_cd, rate, pitch))
                    
                    with open(mp3_path, "wb") as f: f.write(audio_b)
                    with open(srt_path, "w", encoding="utf-8") as f: f.write(srt_b)
                    
                    clean_txt = strip_director_tags(content)
                    with open(txt_path, "w", encoding="utf-8") as f: f.write(clean_txt)
                    
                    with open(hash_path, "w", encoding="utf-8") as f: f.write(chunk_hash)
                    
                    vid_b = None
                    if exp_mp4:
                        progress_bar.progress(idx / len(chapters), text=f"Video: {title} ({idx+1}/{len(chapters)})")
                        try: 
                            vid_b = generate_mp4_video(audio_b, srt_b, dur, progress_bar)
                            with open(mp4_path, "wb") as f: f.write(vid_b)
                        except Exception as e: 
                            st.warning(f"Video fail: {str(e)}")

                    chapter_outputs.append({"title": title, "audio_bytes": audio_b, "srt_text": srt_b, "script_text": content, "video_bytes": vid_b})
                
                completed_count += 1
                update_project_metadata(workspace_dir, doc_name_clean, lang_lbl, f"{gender} - {persona}", len(chapters), completed_count)

            progress_bar.progress(1.0, text="Finalizing Archive...")
            zip_bytes = create_zip_archive(chapter_outputs, full_docx)
            progress_bar.empty()
            st.success("✅ Generation complete!")
            
        except Exception as e:
            st.error(f"Failed: {str(e)}")
            return
        
        st.session_state.audio_history.insert(0, {
            "filename": f"Audiobook_{lang_cd}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            "voice": f"{gender} - {persona}",
            "size": f"{len(zip_bytes) / 1024 / 1024:.2f} MB",
            "zip_bytes": zip_bytes,
            "chapters_count": completed_count,
            "contains_video": exp_mp4,
            "time": datetime.now().strftime("%I:%M %p")
        })

def render_history():
    if not st.session_state.audio_history: return
    st.divider()
    st.subheader("🎧 Current Session Audiobooks")
    
    if len(st.session_state.audio_history) > 1:
        master_zip = create_master_zip_archive(st.session_state.audio_history)
        st.download_button(
            label="📦 Download ALL Audiobooks (Master ZIP)",
            data=master_zip,
            file_name=f"Master_Audiobook_Archive_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            mime="application/zip",
            type="primary",
            width="stretch"
        )
        st.markdown("---")
        
    for idx, track in enumerate(st.session_state.audio_history):
        with st.container():
            vid_stat = "🎬 Video Included" if track.get('contains_video') else "🎧 Audio Only"
            st.markdown(f"**{track['time']}** | 🗣️ {track['voice']} | 📦 Chapters: {track['chapters_count']} | {vid_stat} | 💾 {track['size']}")
            st.download_button(
                label=f"📦 Download Audiobook {idx + 1} (ZIP)",
                data=track['zip_bytes'],
                file_name=track['filename'],
                mime="application/zip",
                width="stretch",
                key=f"zip_dl_{idx}"
            )
            st.markdown("---")

def render_database_tab():
    st.header("🗄️ Past Projects Database")
    st.write("Access previously generated audiobooks and physical folders across sessions.")
    st.divider()
    
    workspace_base = os.path.join(os.getcwd(), "audiobook_workspace")
    if not os.path.exists(workspace_base):
        st.info("No projects found yet. Start generating in the Studio tab!")
        return
        
    projects = []
    for d in os.listdir(workspace_base):
        meta_path = os.path.join(workspace_base, d, "metadata.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, "r", encoding="utf-8") as f:
                    projects.append((d, json.load(f)))
            except: pass
            
    if not projects:
        st.info("No saved projects found in the workspace.")
        return
        
    projects.sort(key=lambda x: x[1].get("last_updated", ""), reverse=True)
    
    for folder_name, meta in projects:
        doc_name = meta.get('doc_name', folder_name)
        comp = meta.get('completed_chapters', 0)
        tot = meta.get('total_chapters', '?')
        status = "✅ Complete" if str(comp) == str(tot) else "⏳ In Progress"
        
        with st.expander(f"{status} | 📁 {doc_name} | {meta.get('language', '')} ({meta.get('voice', '')})"):
            st.write(f"**Last Updated:** {meta.get('last_updated', '')}")
            st.write(f"**Progress:** {comp} / {tot} Chapters Rendered")
            
            workspace_dir = os.path.join(workspace_base, folder_name)
            
            if st.button("📝 Load Project to Studio", key=f"load_db_{folder_name}", type="primary", width="stretch"):
                script_path = os.path.join(workspace_dir, "master_script.txt")
                if os.path.exists(script_path):
                    with open(script_path, "r", encoding="utf-8") as f:
                        st.session_state.pasted_text = f.read()
                        st.session_state.source_doc_name = doc_name
                    st.success("✅ Project loaded! Go to the '🎙️ Audio Studio' tab and select '✍️ Paste Text' to resume generating.")
                else:
                    st.error("No master script found for this project.")
            
            st.markdown("<br>", unsafe_allow_html=True)
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📁 Open Physical Folder", key=f"open_db_{folder_name}", width="stretch"):
                    try:
                        if platform.system() == "Windows": os.startfile(workspace_dir)
                        elif platform.system() == "Darwin": subprocess.Popen(["open", workspace_dir])
                        else: subprocess.Popen(["xdg-open", workspace_dir])
                    except Exception as e: st.error(f"Error: {e}")
            with col2:
                if st.button("📦 Compile & Download ZIP", key=f"zip_db_{folder_name}", width="stretch"):
                    with st.spinner("Compiling Archive..."):
                        zip_data = package_folder_to_zip(workspace_dir)
                        st.download_button(
                            label="⬇️ Click Here to Save ZIP",
                            data=zip_data,
                            file_name=f"Archive_{folder_name}.zip",
                            mime="application/zip",
                            width="stretch",
                            key=f"dl_btn_{folder_name}"
                        )


# ==============================================================================
# 9. APP ENTRY POINT
# ==============================================================================

def main():
    st.set_page_config(page_title="Enterprise Audiobook Studio", page_icon="📚", layout="wide")
    initialize_session_state()
    st.title("📚 Enterprise Multilingual Audiobook Studio")
    st.markdown("Convert massive PDF books or Text scripts into chapter-segmented audiobooks.")
    
    render_error_dashboard()
    lang_cd, lang_lbl, voice_cd, rate, pitch, exp_mp4, gender, persona, sec_voice_cd = render_sidebar()
    
    tab_studio, tab_database = st.tabs(["🎙️ Audio Studio", "🗄️ Projects Database"])
    
    with tab_studio:
        method = st.radio("Input Method:", ["📄 Attach File", "✍️ Paste Text"], horizontal=True, label_visibility="collapsed")
        active_text = ""

        if method == "📄 Attach File":
            doc_lang = st.selectbox("OCR Source Language:", list(LANGUAGES.keys()), index=1)
            files = st.file_uploader("Upload File(s)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
            
            if files:
                engine, gem_key = "", ""
                pdf_files = [f for f in files if f.name.lower().endswith('.pdf')]
                is_single_pdf = (len(files) == 1 and len(pdf_files) == 1)
                
                st.session_state.source_doc_name = "_".join([f.name for f in files])[:50]
                
                if pdf_files:
                    engine = st.radio("PDF Engine:", ["🤖 Gemini Vision AI", "⚡ Local EasyOCR (CPU/GPU)"], index=0)
                    if "Gemini" in engine:
                        default_api_key = os.environ.get("GEMINI_API_KEY", "")
                        gem_key = st.text_input("Gemini API Key:", value=default_api_key, type="password", placeholder="Enter your Google Gemini API Key here")
                
                start_page, end_page = None, None
                if is_single_pdf:
                    doc_temp = fitz.open(stream=files[0].getvalue(), filetype="pdf")
                    st.info(f"📄 **{files[0].name}** ({len(doc_temp)} pages).")
                    col_p1, col_p2 = st.columns(2)
                    with col_p1: start_page = st.number_input("Start Page", min_value=1, max_value=len(doc_temp), value=1)
                    with col_p2: end_page = st.number_input("End Page", min_value=1, max_value=len(doc_temp), value=min(10, len(doc_temp)))
                    doc_temp.close()
                    
                if st.button("🚀 Extract Script", type="primary", width="stretch"):
                    prog = st.progress(0.0, text="Extracting...")
                    all_extracted = []
                    try:
                        for idx, file in enumerate(files):
                            file_bytes = file.read()
                            ext = file.name.lower().split('.')[-1]
                            
                            if ext == "pdf":
                                doc = fitz.open(stream=file_bytes, filetype="pdf")
                                s_page = start_page if is_single_pdf else 1
                                e_page = end_page if is_single_pdf else len(doc)
                                
                                if "Gemini" in engine:
                                    if not gem_key.strip():
                                        raise ValueError("Gemini API key is required when using the Gemini Vision AI engine.")
                                    ext_text = extract_text_with_vision_ai(file_bytes, gem_key, s_page, e_page, prog, file_index=idx, total_files=len(files))
                                else:
                                    ext_text = extract_text_ocr_local(file_bytes, LANGUAGES[doc_lang], s_page, e_page, prog)
                                doc.close()
                            elif ext == "docx": ext_text = clean_extracted_text(extract_text_from_docx(file_bytes))
                            else: ext_text = clean_extracted_text(extract_text_from_txt(file_bytes))
                                
                            if len(files) > 1: all_extracted.append(f"\n\n--- Start of {file.name} ---\n\n" + ext_text)
                            else: all_extracted.append(ext_text)
                        
                        st.session_state.pdf_extracted_text = "\n\n".join(all_extracted)
                        prog.empty()
                        st.success("✅ Extraction Successful!")
                    except Exception as e: 
                        prog.empty()
                        st.error(f"Error during extraction: {str(e)}")
                    
                if st.session_state.pdf_extracted_text:
                    if pdf_files:
                        st.markdown("---")
                        st.markdown("### 🔍 Verification Dashboard")
                        col_pdf, col_text = st.columns([1, 1], gap="large")
                        
                        with col_pdf:
                            st.markdown("#### 📄 Original PDF Viewer")
                            
                            if len(pdf_files) > 1:
                                selected_pdf_name = st.selectbox("Select PDF to view:", [f.name for f in pdf_files])
                                selected_pdf = next(f for f in pdf_files if f.name == selected_pdf_name)
                            else:
                                selected_pdf = pdf_files[0]
                                
                            doc_view = fitz.open(stream=selected_pdf.getvalue(), filetype="pdf") 
                            
                            v_start = start_page if (is_single_pdf and start_page is not None) else 1
                            v_end = end_page if (is_single_pdf and end_page is not None) else len(doc_view)
                            v_start = max(1, min(v_start, len(doc_view)))
                            v_end = max(v_start, min(v_end, len(doc_view)))
                            
                            if v_start < v_end:
                                viewer_page = st.slider("Select Page to View", min_value=v_start, max_value=v_end, value=v_start, key="pdf_viewer_slider")
                            else:
                                viewer_page = v_start
                                st.caption(f"Viewing Page {viewer_page}")

                            pix = doc_view[viewer_page - 1].get_pixmap(dpi=150)
                            st.image(pix.tobytes("jpeg", 90))
                            doc_view.close()
                            
                        with col_text:
                            st.markdown("#### ✍️ Extracted Script Editor")
                            active_text = render_script_tools("pdf_extracted_text", lang_cd, lang_lbl)
                    else:
                        active_text = render_script_tools("pdf_extracted_text", lang_cd, lang_lbl)
        else:
            if st.session_state.source_doc_name == "_".join([f.name for f in files] if 'files' in locals() and files else []): 
                st.session_state.source_doc_name = "Pasted_Text"
            st.session_state.pasted_text = st.text_area("Paste script:", value=st.session_state.pasted_text, height=250)
            if st.session_state.pasted_text.strip():
                active_text = render_script_tools("pasted_text", lang_cd, lang_lbl)

        handle_generation(active_text, lang_cd, lang_lbl, voice_cd, rate, pitch, exp_mp4, gender, persona, sec_voice_cd)
        render_history()
        
    with tab_database:
        render_database_tab()

if __name__ == "__main__":
    main()